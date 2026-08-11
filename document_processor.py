from io import BytesIO
from pathlib import Path

from docx import Document as WordDocument
from pypdf import PdfReader


def read_txt(file_bytes):
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


def read_pdf(file_bytes):
    reader = PdfReader(BytesIO(file_bytes))

    pages = []

    for page in reader.pages:
        page_text = page.extract_text()

        if page_text:
            pages.append(page_text)

    if not pages:
        raise ValueError(
            "No readable text was found in the PDF. "
            "Scanned PDFs require OCR support."
        )

    return "\n\n".join(pages)


def read_docx(file_bytes):
    document = WordDocument(BytesIO(file_bytes))

    contents = []

    # Normal paragraphs
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            contents.append(text)

    # Text contained inside tables
    for table in document.tables:
        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if cells:
                contents.append(" | ".join(cells))

    if not contents:
        raise ValueError(
            "No readable text was found in the DOCX file."
        )

    return "\n\n".join(contents)


def clean_text(text):
    lines = [
        line.strip()
        for line in text.splitlines()
        if line.strip()
    ]

    return "\n".join(lines)


def split_into_chunks(
    text,
    chunk_size=900,
    overlap=150,
):
    chunks = []
    start = 0

    while start < len(text):
        end = min(start + chunk_size, len(text))

        chunk = text[start:end].strip()

        if chunk:
            chunks.append(chunk)

        if end >= len(text):
            break

        start = end - overlap

    return chunks


def process_document(file_name, file_bytes):
    extension = Path(file_name).suffix.lower()

    if extension == ".txt":
        text = read_txt(file_bytes)

    elif extension == ".pdf":
        text = read_pdf(file_bytes)

    elif extension == ".docx":
        text = read_docx(file_bytes)

    else:
        raise ValueError(
            "Unsupported file type. "
            "Only PDF, DOCX and TXT files are supported."
        )

    text = clean_text(text)

    if not text:
        raise ValueError(
            "No readable text was found in the document."
        )

    return split_into_chunks(text)